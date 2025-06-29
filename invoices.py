import os
import json
import re
from datetime import datetime
from typing import Dict, List, Tuple, Optional

try:
    import pdfplumber
    print("pdfplumber imported successfully")
except ImportError:
    print("pdfplumber not found. Install with: pip install pdfplumber")
    exit(1)

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill
    print("openpyxl imported successfully")
except ImportError:
    print("openpyxl not found. Install with: pip install openpyxl")
    exit(1)
    
# kmk, na boa?
try:
    import cv2
    import numpy as np
    from PIL import Image
    import pytesseract
    import fitz  # PyMuPDF
    ADVANCED_OCR_AVAILABLE = True
    print("Advanced OCR libraries available")
except ImportError as e:
    ADVANCED_OCR_AVAILABLE = False
    print(f"Advanced OCR not available: {e}")
    print("   The extractor will work only with PDFs that already contain text")

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    WORD_AVAILABLE = True
    print("python-docx available")
except ImportError:
    WORD_AVAILABLE = False
    print("python-docx not available. Word files will not be created")

class BookExtractor:
    def __init__(self, directory: str = 'pdf_books'):
        self.directory = directory
        self.extracted_books = []
        
    def clean_musical_notation(self, text: str) -> str:
        """Remove musical notation patterns from text"""
        print("Removing musical notations...")
        
        musical_patterns = [
            r'[:]+[-—.:]+[|]*', 
            r'[smtdfrl]+[,\']*\s*[:—.-]+\s*[smtdfrl,\']*',
            r'[|]+\s*[:.—-]+\s*[|]*',
            r'^[sd,\':\-—|.\s]+$',
            r'[drmfslti]+[,\']*\s*[:\-—.]+\s*[drmfslti,\']*', 
            r'\b[a-z][,\']+\s*[:\-—.]+', 
            r'[:\-—.]{3,}', 
            r'^[|]+.*[|]+$', 
            r'^\s*[:.—-]+\s*$', 
            r'\b[smtdfrl]+[,\']*[:\-—.]+[smtdfrl,\']*\b',  
        ]
        
        lines = text.split('\n')
        cleaned_lines = []
        removed_count = 0
        
        for line in lines:
            original_line = line.strip()
            
            if not original_line:
                continue
            
            # Check if it's predominantly musical notation
            musical_chars = len(re.findall(r'[:\-—.|,\'smdftlri]', original_line))
            total_chars = len(original_line.replace(' ', ''))
            
            # If more than 60% are musical symbols, skip
            if total_chars > 0 and (musical_chars / total_chars) > 0.6:
                removed_count += 1
                continue
            
            # Apply pattern cleaning
            cleaned_line = original_line
            for pattern in musical_patterns:
                cleaned_line = re.sub(pattern, '', cleaned_line, flags=re.IGNORECASE)
            
            cleaned_line = ' '.join(cleaned_line.split())
            
            # Check if it still has valid content
            if cleaned_line and len(cleaned_line.strip()) > 2:
                if re.search(r'[a-zA-Z]{3,}', cleaned_line):
                    cleaned_lines.append(cleaned_line)
        
        print(f"   Removed {removed_count} lines of musical notation")
        return '\n'.join(cleaned_lines)

    def preprocess_image_for_ocr(self, image_array: np.ndarray) -> np.ndarray:
        """Preprocess image for better OCR results"""
        if not ADVANCED_OCR_AVAILABLE:
            return image_array
            
        # Convert to grayscale
        if len(image_array.shape) == 3:
            gray = cv2.cvtColor(image_array, cv2.COLOR_RGB2GRAY)
        else:
            gray = image_array
            
        # Bilateral filter
        bilateral = cv2.bilateralFilter(gray, 9, 75, 75)
        
        # Adaptive threshold
        thresh = cv2.adaptiveThreshold(bilateral, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
            cv2.THRESH_BINARY, 11, 2)
        
        # Morphological operations
        kernel = np.ones((1,1), np.uint8)
        opening = cv2.morphologyEx(thresh, cv2.MORPH_OPEN, kernel)
        closing = cv2.morphologyEx(opening, cv2.MORPH_CLOSE, kernel)
        
        # Denoising
        denoised = cv2.fastNlMeansDenoising(closing)
        
        return denoised

    def extract_text_basic(self, pdf_path: str) -> str:
        """Extract text using pdfplumber"""
        print("Extracting text with pdfplumber...")
        all_text = []
        
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    print(f"   Processing page {i+1}/{len(pdf.pages)}")
                    text = page.extract_text()
                    if text and text.strip():
                        cleaned_text = self.clean_musical_notation(text)
                        if cleaned_text.strip():
                            all_text.append(cleaned_text)
        except Exception as e:
            print(f"Error with pdfplumber: {e}")
            
        return '\n'.join(all_text)

    def extract_text_advanced(self, pdf_path: str) -> str:
        if not ADVANCED_OCR_AVAILABLE:
            return self.extract_text_basic(pdf_path)
            
        print("Extracting text with advanced OCR...")
        all_text = []
        
        # First try basic method
        basic_text = self.extract_text_basic(pdf_path)
        if basic_text and len(basic_text.strip()) > 100:
            return basic_text
        
        print("   Insufficient text, applying OCR...")
        try:
            doc = fitz.open(pdf_path)
            for page_num in range(doc.page_count):
                print(f"   OCR page {page_num+1}/{doc.page_count}")
                page = doc[page_num]
                
                # Extract as image
                mat = fitz.Matrix(2.0, 2.0)
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                
                # Convert to numpy
                img_array = np.frombuffer(img_data, dtype=np.uint8)
                img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)
                
                # Preprocess
                processed_img = self.preprocess_image_for_ocr(img)
                
                # OCR
                config = r'--oem 3 --psm 6'
                text = pytesseract.image_to_string(processed_img, config=config, lang='por+eng')
                
                if text.strip():
                    cleaned_text = self.clean_musical_notation(text)
                    if cleaned_text.strip():
                        all_text.append(cleaned_text)
                        
            doc.close()
        except Exception as e:
            print(f"Error with OCR: {e}")
        
        return '\n'.join(all_text)

    def detect_structure(self, text: str) -> Dict:
        print("Detecting text structure...")
        
        lines = text.split('\n')
        structure = {
            'title': '',
            'songs': [],
            'metadata': {
                'total_lines': len([l for l in lines if l.strip()]),
                'extraction_date': datetime.now().isoformat()
            }
        }
        
        current_song = None
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            
            # Detect song/hymn numbers with titles (e.g., "1 O titulo")
            song_match = re.match(r'^(\d+)\s+(.+)', line)
            if song_match and len(song_match.group(2)) > 3: 
                is_title = True
                if i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    if re.match(r'^\d+\s', next_line):
                        numbered_sequence = 0
                        for j in range(i, min(i + 5, len(lines))):
                            if re.match(r'^\d+\s', lines[j].strip()):
                                numbered_sequence += 1
                        if numbered_sequence >= 3:  # 3 ou mais linhas numeradas = estrofes
                            is_title = False
                
                if is_title:
                    # Salvar hino anterior
                    if current_song:
                        structure['songs'].append(current_song)
                    
                    # Criar novo hino
                    current_song = {
                        'number': song_match.group(1),
                        'title': song_match.group(2).strip(),
                        'verses': [],
                        'chorus': []
                    }
                    i += 1
                    continue
            
            # Detect main title (usually uppercase and long)
            if line.isupper() and len(line) > 10 and not structure['title'] and not current_song:
                structure['title'] = line
                i += 1
                continue
            
            # Detect chorus/refrain
            chorus_keywords = ['CHORUS', 'CORO', 'REFRAIN', 'CÔRO']
            if any(keyword in line.upper() for keyword in chorus_keywords):
                if current_song:
                    # Coletar linhas do coro
                    chorus_lines = []
                    j = i + 1
                    while j < len(lines) and j < i + 30: 
                        chorus_line = lines[j].strip()
                        if not chorus_line:
                            j += 1
                            continue
                        
                        if re.match(r'^\d+\s', chorus_line):
                            if len(chorus_line.split()) > 3:
                                break
                            break
                        
                        # Parar se encontrar outro coro
                        if any(keyword in chorus_line.upper() for keyword in chorus_keywords):
                            break
                        
                        chorus_lines.append(chorus_line)
                        j += 1
                    
                    current_song['chorus'] = chorus_lines
                    i = j
                    continue
            
            # Detect numbered verses (melhorado)
            verse_match = re.match(r'^(\d+)\s+(.+)', line)
            if verse_match and current_song:
                verse_num = verse_match.group(1)
                verse_text = verse_match.group(2)
                
                if len(verse_text.split()) <= 15:
                    verse_lines = [verse_text]
                    j = i + 1
                    
                    while j < len(lines) and len(verse_lines) < 6:  # Máximo 6 linhas por estrofe
                        next_line = lines[j].strip()
                        if not next_line:
                            j += 1
                            continue
                        
                        # Parar se encontrar nova estrofe numerada
                        if re.match(r'^\d+\s', next_line):
                            break
                        
                        # Parar se encontrar coro
                        if any(keyword in next_line.upper() for keyword in chorus_keywords):
                            break
                        
                        verse_lines.append(next_line)
                        j += 1
                    
                    # Adicionar estrofe apenas se tiver conteúdo válido
                    if len(' '.join(verse_lines).strip()) > 10:
                        current_song['verses'].append({
                            'number': verse_num,
                            'lines': verse_lines
                        })
                    
                    i = j
                    continue
            
            if current_song and len(line) > 5 and not re.match(r'^\d+', line):
                # Verificar se não é início de novo hino
                words = line.split()
                if len(words) > 1 and not line.isupper():
                    # Adicionar como estrofe sem número
                    verse_lines = [line]
                    j = i + 1
                    
                    while j < len(lines) and len(verse_lines) < 4:
                        next_line = lines[j].strip()
                        if not next_line:
                            j += 1
                            continue
                        
                        if re.match(r'^\d+\s', next_line) and len(next_line.split()) > 3:
                            break
                        
                        if any(keyword in next_line.upper() for keyword in chorus_keywords):
                            break
                        
                        verse_lines.append(next_line)
                        j += 1
                    
                    if len(current_song['verses']) == 0 or len(' '.join(verse_lines)) > 15:
                        current_song['verses'].append({
                            'number': str(len(current_song['verses']) + 1),
                            'lines': verse_lines
                        })
                    
                    i = j
                    continue
            
            i += 1
        
        # Add last song
        if current_song:
            structure['songs'].append(current_song)
        
        # Filtrar hinos com conteúdo válido
        valid_songs = []
        for song in structure['songs']:
            if song.get('verses') or song.get('chorus'):
                valid_songs.append(song)
        
        structure['songs'] = valid_songs
        
        print(f"   Found {len(structure['songs'])} songs/hymns")
        for song in structure['songs'][:3]:  # Show first 3 for debugging
            print(f"     {song.get('number', 'N/A')}. {song.get('title', 'No title')[:30]}... - {len(song.get('verses', []))} verses, {len(song.get('chorus', []))} chorus lines")
        
        return structure

    def create_excel_file(self, book_data: Dict, filename: str):
        """Create Excel file with extracted data"""
        print(f"Creating Excel file: {filename}")
        
        wb = Workbook()
        ws = wb.active
        ws.title = 'Extracted Hymnal'
        
        # Headers
        headers = ['Number', 'Title', 'Type', 'Verse', 'Content']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True)
        
        row = 2
        
        # Book title
        if book_data.get('title'):
            ws.cell(row=row, column=3, value='MAIN TITLE')
            ws.cell(row=row, column=5, value=book_data['title'])
            row += 1
        
        # Songs
        for song in book_data.get('songs', []):
            # Song title
            ws.cell(row=row, column=1, value=song.get('number', ''))
            ws.cell(row=row, column=2, value=song.get('title', ''))
            ws.cell(row=row, column=3, value='TITLE')
            row += 1
            
            # Verses
            for verse in song.get('verses', []):
                ws.cell(row=row, column=1, value=song.get('number', ''))
                ws.cell(row=row, column=3, value='VERSE')
                ws.cell(row=row, column=4, value=verse.get('number', ''))
                ws.cell(row=row, column=5, value=' / '.join(verse.get('lines', [])))
                row += 1
                
            # Chorus
            if song.get('chorus'):
                ws.cell(row=row, column=1, value=song.get('number', ''))
                ws.cell(row=row, column=3, value='CHORUS')
                ws.cell(row=row, column=5, value=' / '.join(song['chorus']))
                row += 1
        
        # Adjust columns
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            ws.column_dimensions[column_letter].width = min(max_length + 2, 80)
        
        wb.save(filename)
        print(f"   Excel saved: {filename}")

    def create_json_file(self, book_data: Dict, filename: str):
        """Create JSON file with extracted data"""
        print(f"Creating JSON file: {filename}")
        
        json_structure = {
            "metadata": {
                "extraction_date": datetime.now().isoformat(),
                "total_songs": len(book_data.get('songs', [])),
                "title": book_data.get('title', ''),
                "extractor_version": "2.1"
            },
            "content": {
                "title": book_data.get('title', ''),
                "songs": []
            }
        }
        
        for song in book_data.get('songs', []):
            song_data = {
                "number": song.get('number', ''),
                "title": song.get('title', ''),
                "verses": {},
                "chorus": song.get('chorus', [])
            }
            
            for i, verse in enumerate(song.get('verses', [])):
                verse_key = f"estrofe_{verse.get('number', i+1)}"
                song_data["verses"][verse_key] = {
                    "number": verse.get('number', str(i+1)),
                    "lines": verse.get('lines', [])
                }
            
            json_structure["content"]["songs"].append(song_data)
        
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(json_structure, f, indent=2, ensure_ascii=False)
        
        print(f"   JSON saved: {filename}")

    def create_word_file(self, book_data: Dict, filename: str):
        """Create Word document with extracted data"""
        if not WORD_AVAILABLE:
            print("Skipping Word creation (library not available)")
            return
            
        print(f"Creating Word file: {filename}")
        
        doc = Document()
        
        # Main title
        if book_data.get('title'):
            title = doc.add_heading(book_data['title'], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Songs
        for song in book_data.get('songs', []):
            # Song title
            song_title = f"{song.get('number', '')}. {song.get('title', '')}"
            doc.add_heading(song_title, level=1)
            
            # Verses
            for verse in song.get('verses', []):
                doc.add_heading(f"Stanza {verse.get('number', '')}", level=2)
                
                for line in verse.get('lines', []):
                    doc.add_paragraph(line)
                
                doc.add_paragraph()  # Space
            
            if song.get('chorus'):
                doc.add_heading("Chorus", level=2)
                for line in song['chorus']:
                    doc.add_paragraph(line)
                
                doc.add_paragraph()
            
            doc.add_page_break()
        
        doc.save(filename)
        print(f"   Word saved: {filename}")

    def process_files(self):
        """Process all PDF files in the directory"""
        print(f"Starting processing in: {self.directory}")
        
        if not os.path.exists(self.directory):
            os.makedirs(self.directory)
            print(f"Directory created: {self.directory}")
            print("   Place your PDF files in this directory and run again.")
            return
        
        files = [f for f in os.listdir(self.directory) if f.lower().endswith('.pdf')]
        if not files:
            print("No PDF files found!")
            print(f"   Place PDF files in: {os.path.abspath(self.directory)}")
            return
        
        print(f"Found {len(files)} PDF file(s)")
        
        for i, file in enumerate(files, 1):
            print(f"\n{'='*60}")
            print(f"Processing [{i}/{len(files)}]: {file}")
            print('='*60)
            
            file_path = os.path.join(self.directory, file)
            
            try:
                if ADVANCED_OCR_AVAILABLE:
                    text = self.extract_text_advanced(file_path)
                else:
                    text = self.extract_text_basic(file_path)
                
                if not text.strip():
                    print("No text extracted")
                    continue
                
                print(f"Text extracted: {len(text)} characters")
                
                book_structure = self.detect_structure(text)
                
                base_name = os.path.splitext(file)[0]
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                
                excel_file = f"{base_name}_extracted_{timestamp}.xlsx"
                json_file = f"{base_name}_extracted_{timestamp}.json"
                word_file = f"{base_name}_extracted_{timestamp}.docx"
                
                self.create_excel_file(book_structure, excel_file)
                self.create_json_file(book_structure, json_file)
                self.create_word_file(book_structure, word_file)
                
                self.extracted_books.append({
                    'original_file': file,
                    'excel_file': excel_file,
                    'json_file': json_file,
                    'word_file': word_file if WORD_AVAILABLE else 'N/A',
                    'songs_found': len(book_structure.get('songs', [])),
                    'status': 'SUCCESS'
                })
                
                print(f"Processing completed!")
                
            except Exception as e:
                print(f"Error: {e}")
                self.extracted_books.append({
                    'original_file': file,
                    'status': f'ERROR: {str(e)}'
                })
        
        self.generate_report()

    def generate_report(self):
        print(f"\n{'='*60}")
        print("FINAL REPORT")
        print('='*60)
        
        successful = [b for b in self.extracted_books if b.get('status') == 'SUCCESS']
        failed = [b for b in self.extracted_books if 'ERROR' in b.get('status', '')]
        
        print(f"Total files: {len(self.extracted_books)}")
        print(f"Successful: {len(successful)}")
        print(f"Failed: {len(failed)}")
        
        if successful:
            total_songs = sum(b.get('songs_found', 0) for b in successful)
            print(f"Total songs/hymns extracted: {total_songs}")
            
            print("\nFiles created:")
            for book in successful:
                print(f"   {book['original_file']}:")
                print(f"     - {book['excel_file']}")
                print(f"     - {book['json_file']}")
                if book.get('word_file') != 'N/A':
                    print(f"     - {book['word_file']}")
                print(f"     - Songs: {book['songs_found']}")
        
        if failed:
            print("\nFiles with errors:")
            for book in failed:
                print(f"   {book['original_file']}: {book['status']}")
        
        report_file = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        with open(report_file, 'w', encoding='utf-8') as f:
            json.dump({
                'timestamp': datetime.now().isoformat(),
                'summary': {
                    'total': len(self.extracted_books),
                    'successful': len(successful),
                    'failed': len(failed)
                },
                'files': self.extracted_books
            }, f, indent=2, ensure_ascii=False)
        
        print(f"\nDetailed report saved in: {report_file}")

def main():
    print("="*60)
    print("PDF HYMNAL AND BOOK EXTRACTOR")
    print("="*60)
    print("Version 2.1 - With improved verse detection")
    print()
    
    print("Checking dependencies...")
    if ADVANCED_OCR_AVAILABLE:
        print("   Advanced OCR: Available")
    else:
        print("   Advanced OCR: Not available")
        print("     (Works only with PDFs that already contain text)")
    
    if WORD_AVAILABLE:
        print("   Word Export: Available")
    else:
        print("   Word Export: Not available")
    
    print()
    
    # Process
    extractor = BookExtractor('pdf_books')
    extractor.process_files()
    
    print("\nProcess finished!")
    print("Check the generated files in the current directory")

if __name__ == '__main__':
    main()