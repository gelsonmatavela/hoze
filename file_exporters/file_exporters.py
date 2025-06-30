"""File export utilities for creating Excel, JSON, and Word files."""

import json
from datetime import datetime
from typing import Dict
from openpyxl import Workbook
from openpyxl.styles import Font
from config import WORD_AVAILABLE

if WORD_AVAILABLE:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH


class FileExporter:
    """Handles exporting extracted data to various file formats."""
    
    def __init__(self):
        self.word_available = WORD_AVAILABLE
    
    def create_excel_file(self, book_data: Dict, filename: str):
        """Create Excel file with extracted data."""
        print(f"Creating Excel file: {filename}")
        
        wb = Workbook()
        ws = wb.active
        ws.title = 'Extracted Hymnal'
        
        # Headers
        headers = ['Number', 'Title', 'Type', 'Verse', 'Content']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True)
        
        row = 2
        
        # Book title
        if book_data.get('title'):
            ws.cell(row=row, column=3, value='MAIN TITLE')
            ws.cell(row=row, column=5, value=book_data['title'])
            row += 1
        
        # Songs
        for song in book_data.get('songs', []):
            row = self._add_song_to_excel(ws, song, row)
        
        # Adjust columns
        self._adjust_excel_columns(ws)
        
        wb.save(filename)
        print(f"   Excel saved: {filename}")
    
    def create_json_file(self, book_data: Dict, filename: str):
        """Create JSON file with extracted data."""
        print(f"Creating JSON file: {filename}")
        
        json_structure = {
            "metadata": {
                "extraction_date": datetime.now().isoformat(),
                "total_songs": len(book_data.get('songs', [])),
                "title": book_data.get('title', ''),
                "extractor_version": "2.1"
            },
            "content": {
                "title": book_data.get('title', ''),
                "songs": []
            }
        }
        
        for song in book_data.get('songs', []):
            song_data = self._format_song_for_json(song)
            json_structure["content"]["songs"].append(song_data)
        
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(json_structure, f, indent=2, ensure_ascii=False)
        
        print(f"   JSON saved: {filename}")
    
    def create_word_file(self, book_data: Dict, filename: str):
        """Create Word document with extracted data."""
        if not self.word_available:
            print("Skipping Word creation (library not available)")
            return
            
        print(f"Creating Word file: {filename}")
        
        doc = Document()
        
        # Main title
        if book_data.get('title'):
            title = doc.add_heading(book_data['title'], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Songs
        for song in book_data.get('songs', []):
            self._add_song_to_word(doc, song)
        
        doc.save(filename)
        print(f"   Word saved: {filename}")
    
    def _add_song_to_excel(self, ws, song: Dict, row: int) -> int:
        """Add a song to the Excel worksheet."""
        # Song title
        ws.cell(row=row, column=1, value=song.get('number', ''))
        ws.cell(row=row, column=2, value=song.get('title', ''))
        ws.cell(row=row, column=3, value='TITLE')
        row += 1
        
        # Verses
        for verse in song.get('verses', []):
            ws.cell(row=row, column=1, value=song.get('number', ''))
            ws.cell(row=row, column=3, value='VERSE')
            ws.cell(row=row, column=4, value=verse.get('number', ''))
            ws.cell(row=row, column=5, value=' / '.join(verse.get('lines', [])))
            row += 1
            
        # Chorus
        if song.get('chorus'):
            ws.cell(row=row, column=1, value=song.get('number', ''))
            ws.cell(row=row, column=3, value='CHORUS')
            ws.cell(row=row, column=5, value=' / '.join(song['chorus']))
            row += 1
        
        return row
    
    def _adjust_excel_columns(self, ws):
        """Adjust column widths in Excel worksheet."""
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            ws.column_dimensions[column_letter].width = min(max_length + 2, 80)
    
    def _format_song_for_json(self, song: Dict) -> Dict:
        """Format a song for JSON export."""
        song_data = {
            "number": song.get('number', ''),
            "title": song.get('title', ''),
            "verses": {},
            "chorus": song.get('chorus', [])
        }
        
        for i, verse in enumerate(song.get('verses', [])):
            verse_key = f"estrofe_{verse.get('number', i+1)}"
            song_data["verses"][verse_key] = {
                "number": verse.get('number', str(i+1)),
                "lines": verse.get('lines', [])
            }
        
        return song_data
    
    def _add_song_to_word(self, doc, song: Dict):
        """Add a song to the Word document."""
        # Song title
        song_title = f"{song.get('number', '')}. {song.get('title', '')}"
        doc.add_heading(song_title, level=1)
        
        # Verses
        for verse in song.get('verses', []):
            doc.add_heading(f"Stanza {verse.get('number', '')}", level=2)
            
            for line in verse.get('lines', []):
                doc.add_paragraph(line)
            
            doc.add_paragraph()  # Space
        
        # Chorus
        if song.get('chorus'):
            doc.add_heading("Chorus", level=2)
            for line in song['chorus']:
                doc.add_paragraph(line)
            
            doc.add_paragraph()
        
        doc.add_page_break()